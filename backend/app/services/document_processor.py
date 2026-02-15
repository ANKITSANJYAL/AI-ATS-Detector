"""
Document processor service.
Handles document parsing, text extraction, and preprocessing.
"""
import io

import pypdf
from docx import Document

from app.core.logging import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """
    Document processing service.
    Extracts text from PDF and DOCX files with structure preservation.
    """

    async def extract_text(
        self,
        file_content: bytes,
        mime_type: str
    ) -> str:
        """
        Extract text from document.

        Args:
            file_content: Document binary content
            mime_type: Document MIME type

        Returns:
            Extracted text content

        Raises:
            ValueError: If unsupported document type
        """
        if mime_type == "application/pdf":
            return await self._extract_pdf(file_content)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return await self._extract_docx(file_content)
        elif mime_type == "text/plain":
            return file_content.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported document type: {mime_type}")

    async def extract_structured_text(
        self,
        file_content: bytes,
        mime_type: str
    ) -> list[dict]:
        """
        Extract text from document with structure metadata.

        Args:
            file_content: Document binary content
            mime_type: Document MIME type

        Returns:
            List of text blocks with formatting metadata

        Raises:
            ValueError: If unsupported document type
        """
        if mime_type == "application/pdf":
            return await self._extract_pdf_structured(file_content)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return await self._extract_docx_structured(file_content)
        elif mime_type == "text/plain":
            # For plain text, split into paragraphs
            text = file_content.decode('utf-8', errors='ignore')
            blocks = []
            for paragraph in text.split('\n\n'):
                para = paragraph.strip()
                if not para:
                    continue
                # Detect heading-like paragraphs: short, no trailing
                # sentence-ending punctuation, or ALL-CAPS
                is_heading = (
                    len(para) < 80
                    and not para.endswith(('.', '!', '?', ',', ';'))
                ) or para.isupper()
                blocks.append({
                    'text': para,
                    'type': 'heading' if is_heading else 'paragraph',
                    'font_size': 'large' if is_heading else 'normal',
                    'bold': is_heading,
                })
            return blocks if blocks else [{'text': text, 'type': 'paragraph', 'font_size': 'normal', 'bold': False}]
        else:
            raise ValueError(f"Unsupported document type: {mime_type}")

    async def _extract_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF.

        Args:
            file_content: PDF binary content

        Returns:
            Extracted text
        """
        try:
            pdf_file = io.BytesIO(file_content)
            reader = pypdf.PdfReader(pdf_file)

            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            full_text = "\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from PDF")
            return full_text

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise ValueError(f"Failed to extract PDF content: {str(e)}") from e

    async def _extract_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX.

        Args:
            file_content: DOCX binary content

        Returns:
            Extracted text
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            full_text = "\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Failed to extract DOCX content: {str(e)}") from e

    async def _extract_pdf_structured(self, file_content: bytes) -> list[dict]:
        """
        Extract structured text from PDF with formatting metadata.

        Args:
            file_content: PDF binary content

        Returns:
            List of text blocks with metadata
        """
        try:
            pdf_file = io.BytesIO(file_content)
            reader = pypdf.PdfReader(pdf_file)

            blocks = []
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    # Split into paragraphs (double newline or significant spacing)
                    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

                    for para in paragraphs:
                        # Detect if it looks like a heading (short, all caps, or ends without period)
                        is_heading = (
                            len(para) < 60 and
                            (para.isupper() or not para.endswith(('.', '!', '?', ',')))
                        )

                        blocks.append({
                            'text': para,
                            'type': 'heading' if is_heading else 'paragraph',
                            'page': page_num + 1,
                            'font_size': 'large' if is_heading else 'normal',
                            'bold': is_heading,
                        })

            logger.info(f"Extracted {len(blocks)} structured blocks from PDF")
            return blocks

        except Exception as e:
            logger.error(f"PDF structured extraction failed: {e}")
            raise ValueError(f"Failed to extract PDF structure: {str(e)}") from e

    async def _extract_docx_structured(self, file_content: bytes) -> list[dict]:
        """
        Extract structured text from DOCX with formatting metadata.

        Args:
            file_content: DOCX binary content

        Returns:
            List of text blocks with metadata
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            blocks = []
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue

                # Detect paragraph type
                style_name = para.style.name.lower() if para.style else ''
                is_heading = 'heading' in style_name or 'title' in style_name

                # Get formatting
                font_size = 'normal'
                if para.runs:
                    # Check first run for font size
                    first_run = para.runs[0]
                    if first_run.font.size:
                        size_pt = first_run.font.size.pt
                        if size_pt > 14:
                            font_size = 'large'
                        elif size_pt > 11:
                            font_size = 'medium'

                # Check if bold
                is_bold = False
                if para.runs:
                    # If majority of runs are bold
                    bold_chars = sum(len(r.text) for r in para.runs if r.bold)
                    total_chars = sum(len(r.text) for r in para.runs)
                    is_bold = bold_chars > total_chars / 2 if total_chars > 0 else False

                blocks.append({
                    'text': para.text.strip(),
                    'type': 'heading' if is_heading else 'paragraph',
                    'font_size': font_size,
                    'bold': is_bold or is_heading,
                    'style': style_name,
                })

            logger.info(f"Extracted {len(blocks)} structured blocks from DOCX")
            return blocks

        except Exception as e:
            logger.error(f"DOCX structured extraction failed: {e}")
            raise ValueError(f"Failed to extract DOCX structure: {str(e)}") from e

    def preprocess_text(self, text: str) -> str:
        """
        Preprocess extracted text.
        Cleans and normalizes text for analysis.

        Args:
            text: Raw extracted text

        Returns:
            Preprocessed text
        """
        # Remove excessive whitespace
        text = " ".join(text.split())

        # Remove common artifacts
        text = text.replace("\x00", "")

        return text.strip()


# Global document processor instance
_document_processor: DocumentProcessor | None = None


def get_document_processor() -> DocumentProcessor:
    """
    Get document processor instance.

    Returns:
        Document processor instance
    """
    global _document_processor

    if _document_processor is None:
        _document_processor = DocumentProcessor()
        logger.info("Document processor initialized")

    return _document_processor
